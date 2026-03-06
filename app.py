from flask import Flask, request, render_template, send_file, jsonify, redirect, url_for
from anthropic import Anthropic
import docx, fitz, os, uuid, json, sqlite3, re
from datetime import datetime
from pathlib import Path
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB
app.secret_key = 'super-secret-key-for-flashing-messages'

# Handle missing API key gracefully for development
anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
ANTHROPIC_CLIENT = Anthropic(api_key=anthropic_key) if anthropic_key else None

ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.txt'}

# --- Database Setup ---
DB_PATH = 'jobs.db'

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS jobs (
            id TEXT PRIMARY KEY,
            filename TEXT,
            status TEXT,
            style TEXT,
            compliance_score REAL,
            total_changes INTEGER,
            error_message TEXT,
            created_at TEXT
        )
    ''')
    conn.commit()
    conn.close()

# --- Utility Functions ---
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def allowed_file(filename):
    return '.' in filename and \
           os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS

def save_job(job_id, filename, style):
    conn = get_db()
    conn.execute(
        'INSERT INTO jobs (id, filename, status, style, created_at) VALUES (?, ?, ?, ?, ?)',
        (job_id, filename, 'processing', style, datetime.now().isoformat())
    )
    conn.commit()
    conn.close()

def update_job(job_id, status, score=None, changes=None, error=None):
    conn = get_db()
    conn.execute(
        'UPDATE jobs SET status = ?, compliance_score = ?, total_changes = ?, error_message = ? WHERE id = ?',
        (status, score, changes, error, job_id)
    )
    conn.commit()
    conn.close()

def get_job(job_id):
    conn = get_db()
    job = conn.execute('SELECT * FROM jobs WHERE id = ?', (job_id,)).fetchone()
    conn.close()
    return dict(job) if job else None

# --- Document Processing Engine ---

# --- In-Memory Job Store (Demo Only) ---
JOBS_IR = {}

# --- Document Processing Engine ---

def ingest_document(filepath, job_id):
    """Router for different file types into a structured IR dict."""
    ext = os.path.splitext(filepath)[1].lower()
    
    ir = {
        'job_id': job_id,
        'source_format': ext[1:],
        'paragraphs': [],
        'tables': [],
        'raw_text': '',
        'word_count': 0,
        'title': None,
        'abstract': None,
        'sections': [],
        'citations_raw': [],
        'references_raw': [],
        'detected_style': None
    }

    if ext == '.docx':
        ir = ingest_docx(filepath, ir)
    elif ext == '.pdf':
        ir = ingest_pdf(filepath, ir)
    elif ext == '.txt':
        ir = ingest_txt(filepath, ir)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

    # Common metrics
    ir['raw_text'] = ' '.join([p['text'] for p in ir['paragraphs']])
    ir['word_count'] = len(ir['raw_text'].split())
    
    return ir

def ingest_docx(filepath, ir):
    """Extracts structured content from .docx files including formatting metadata."""
    doc = docx.Document(filepath)
    
    # 1. Extract Paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        bold = any(run.bold for run in para.runs)
        font_size = None
        if para.runs and para.runs[0].font.size:
            font_size = para.runs[0].font.size.pt
            
        ir['paragraphs'].append({
            'id': f'p_{i}',
            'text': text,
            'style': para.style.name,
            'is_bold': bold,
            'font_size': font_size,
            'type': 'unknown',
            'heading_level': 0
        })

    # 2. Extract Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        ir['tables'].append({'caption': '', 'rows': rows})
        
    return ir

def ingest_pdf(filepath, ir):
    """Extracts text from PDF files using PyMuPDF (fitz) with reading order sorting."""
    doc = fitz.open(filepath)
    all_blocks = []
    
    for page_num, page in enumerate(doc):
        # Extract blocks: (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text('blocks')
        for b in blocks:
            if b[6] == 0: # text block
                text = b[4].strip()
                if text:
                    all_blocks.append({
                        'page': page_num,
                        'y0': b[1],
                        'text': text
                    })
    
    # Sort blocks by page and then y-coordinate (reading order)
    all_blocks.sort(key=lambda x: (x['page'], x['y0']))
    
    for i, block in enumerate(all_blocks):
        text = block['text']
        
        # Estimate headings based on properties
        is_heading_candidate = len(text) < 100 and (text.isupper() or (text[-1].isalnum() and not text.endswith('.')))
        
        ir['paragraphs'].append({
            'id': f'p_{i}',
            'text': text,
            'style': 'Normal' if not is_heading_candidate else 'Heading Estimated',
            'is_bold': False, # fitz requires more logic for bold detection
            'font_size': None,
            'type': 'unknown',
            'heading_level': 0
        })
        
    return ir

def ingest_txt(filepath, ir):
    """Simple text ingestion splitting by double newline."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    paras = content.split('\n\n')
    for i, text in enumerate(paras):
        if text.strip():
            ir['paragraphs'].append({
                'id': f'p_{i}',
                'text': text.strip(),
                'style': 'Normal',
                'is_bold': False,
                'font_size': None,
                'type': 'unknown',
                'heading_level': 0
            })
    return ir

def detect_structure(ir):
    """
    Annotates each paragraph with its structural role using heuristics.
    """
    has_title = False
    in_references = False
    
    # Regex Patterns
    APA_PAT = r'\([A-Z][a-z]+(?:(?:\s+&\s+|,\s+)[A-Z][a-z]+)*,?\s+\d{4}(?:,\s+p\.?\s*\d+)?\)'
    IEEE_PAT = r'\[\d+(?:,\s*\d+)*\]'

    apa_count = 0
    num_count = 0

    for i, para in enumerate(ir['paragraphs']):
        text = para['text']
        text_lower = text.lower()
        font_size = para.get('font_size') or 0
        style = para.get('style', '')

        # 1. TITLE detection
        if not has_title:
            if (font_size > 14 or 'Title' in style) or (text.isupper() and len(text) < 150):
                ir['title'] = text
                para['type'] = 'title'
                has_title = True
                continue

        # 2. ABSTRACT detection
        if not ir['abstract']:
            words = text.split()
            if text_lower.startswith('abstract'):
                ir['abstract'] = text
                para['type'] = 'abstract'
                continue
            elif has_title and 100 <= len(words) <= 400 and not any(s['paragraph_id'] == para['id'] for s in ir['sections']):
                # Paragraph after title/before headings that looks like an abstract
                ir['abstract'] = text
                para['type'] = 'abstract'
                continue

        # 3. SECTION HEADINGS
        is_heading = False
        level = 0
        if 'Heading' in style:
            match = re.search(r'\d', style)
            level = int(match.group()) if match else 1
            is_heading = True
        elif (font_size >= 13 and len(text) < 100 and not text.endswith('.')) or (text.isupper() and len(text) < 80):
            level = 1
            is_heading = True

        if is_heading:
            para['type'] = 'heading'
            para['heading_level'] = level
            ir['sections'].append({
                'title': text,
                'level': level,
                'paragraph_id': para['id']
            })
            
            # 4. REFERENCE LIST detection
            if any(kw in text_lower for kw in ['references', 'bibliography', 'works cited']):
                in_references = True
            continue

        # 5. REFERENCE paragraphs
        if in_references:
            para['type'] = 'reference'
            ir['references_raw'].append(text)
            continue

        # 6. CITATION detection (Body scans)
        apa_matches = re.findall(APA_PAT, text)
        ieee_matches = re.findall(IEEE_PAT, text)
        
        ir['citations_raw'].extend(apa_matches)
        ir['citations_raw'].extend(ieee_matches)
        
        apa_count += len(apa_matches)
        num_count += len(ieee_matches)

        # Default to Body
        para['type'] = 'body'

    # DETECTED STYLE
    if apa_count > num_count:
        ir['detected_style'] = 'APA'
    elif num_count > 0:
        ir['detected_style'] = 'Vancouver'
    else:
        ir['detected_style'] = 'Unknown'

    return ir


def extract_rules(style_name: str) -> dict:
    """Return a rules dict for the selected style."""
    RULES = {
      'APA 7th Edition': {
        'citation_format': 'author_year',
        'citation_pattern': '(Author, Year)' ,
        'reference_order': 'alphabetical',
        'heading_style': {'1': 'Bold, Centered', '2': 'Bold, Left', '3': 'Bold Italic, Left'},
        'abstract_max_words': 250,
        'doi_format': 'https://doi.org/...',
        'line_spacing': 'double',
        'font': 'Times New Roman 12pt',
        'required_sections': ['Abstract', 'Introduction', 'Method', 'Results', 'Discussion', 'References'],
      },
      'Vancouver': {
        'citation_format': 'numbered',
        'citation_pattern': '[1]',
        'reference_order': 'order_of_appearance',
        'heading_style': {'1': 'Bold', '2': 'Italic', '3': 'Plain'},
        'abstract_max_words': 300,
        'required_sections': ['Abstract', 'Introduction', 'Methods', 'Results', 'Discussion', 'References'],
      },
      'IEEE': {
        'citation_format': 'numbered',
        'citation_pattern': '[1]',
        'reference_order': 'order_of_appearance',
        'heading_style': {'1': 'Bold, Roman Numeral', '2': 'Italic, Letter'},
        'required_sections': ['Abstract', 'Introduction', 'Methodology', 'Results', 'Conclusion', 'References'],
      },
      'MLA 9th': {
        'citation_format': 'author_page',
        'citation_pattern': '(Author Page)',
        'reference_order': 'alphabetical',
        'heading_style': {'1': 'Bold', '2': 'Bold Italic'},
        'required_sections': ['Works Cited'],
      }
    }
    return RULES.get(style_name, RULES['APA 7th Edition'])

def format_document(ir: dict, style: str) -> dict:
    """
    The core AI function. Orchestrates Citation, Reference, and Heading formatting.
    """
    rules = extract_rules(style)
    change_log = []

    ## STEP 1: Format Citations via Claude
    if ANTHROPIC_CLIENT and ir.get('citations_raw'):
        citation_sample = ir['citations_raw'][:15]  # Process first 15 for demo
        prompt = f'''You are an academic formatting expert.
Reformat these in-text citations to {style} format.
Current citations: {json.dumps(citation_sample)}
Rules: {json.dumps(rules)}
Return a JSON object: {{"formatted": ["new citation 1", ...]}}
Return ONLY the JSON, no explanation.'''

        try:
            response = ANTHROPIC_CLIENT.messages.create(
                model='claude-3-haiku-20240307',
                max_tokens=1000,
                messages=[{'role':'user','content':prompt}]
            )
            raw_text = response.content[0].text
            json_match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                formatted_cites = result.get('formatted', [])
                for orig, new in zip(citation_sample, formatted_cites):
                    if orig != new:
                        change_log.append({'type':'citation','before':orig,'after':new,
                                           'rule':f'{style} in-text citation format'})
                ir['citations_formatted'] = formatted_cites
        except Exception as e:
            print(f"Citation Formatting Error: {e}")

    ## STEP 2: Format Reference List via Claude
    if ANTHROPIC_CLIENT and ir.get('references_raw'):
        refs_text = '\n'.join(ir['references_raw'][:20])  # max 20 refs
        prompt = f'''You are an academic formatting expert.
Reformat these references to {style} format.
References:\n{refs_text}
Rules: citation_format={rules['citation_format']}, order={rules['reference_order']}
Return a JSON object: {{"references": ["formatted ref 1", ...]}}
Sort them correctly per the style. Return ONLY JSON.'''

        try:
            response = ANTHROPIC_CLIENT.messages.create(
                model='claude-3-haiku-20240307',
                max_tokens=2500,
                messages=[{'role':'user','content':prompt}]
            )
            raw_text = response.content[0].text
            json_match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                formatted_refs = result.get('references', [])
                for orig, new in zip(ir['references_raw'], formatted_refs):
                    if orig.strip() != new.strip():
                        change_log.append({'type':'reference','before':orig,'after':new,
                                           'rule':f'{style} reference list format'})
                ir['references_formatted'] = formatted_refs
        except Exception as e:
            print(f"Reference Formatting Error: {e}")

    ## STEP 3: Format Headings (rule-based)
    heading_rule = rules.get('heading_style', {})
    heading_count = 0
    for para in ir['paragraphs']:
        if para['type'] == 'heading':
            level = str(para.get('heading_level', 1))
            style_desc = heading_rule.get(level, 'Bold')
            original = para['text']
            new_text = para['text']
            
            if 'Centered' in style_desc and style == 'APA 7th Edition' and level == '1':
                para['align'] = 'center'
                
            if 'Roman Numeral' in style_desc and style == 'IEEE':
                roman = ['I','II','III','IV','V','VI','VII','VIII','IX','X']
                if not re.match(r'^[IVX]+\.', original):
                    para['text'] = f'{roman[min(heading_count, 9)]}. {original}'
                    new_text = para['text']
            
            heading_count += 1
            if original != new_text:
                change_log.append({'type':'heading','before':original,'after':new_text,
                                   'rule':f'{style} heading {level} format'})

    ir['change_log'] = change_log
    ir['style_applied'] = style
    ir['formatted'] = True
    return ir


def validate_document(ir):
    """Calculates a simulated compliance score for the intermediate representation."""
    return {'score': 0.92, 'issues': ['Manual review recommended for reference links']}

def render_docx(ir: dict, job_id: str) -> str:
    """Renders the IR back into a professionally formatted .docx file."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    output_filename = f"{job_id}_formatted.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
    
    doc = docx.Document()

    # Set document-wide font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    ## Add title
    if ir.get('title'):
        # Title is level 0 in Paperpal UI
        title_para = doc.add_heading(ir['title'], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ## Add abstract
    if ir.get('abstract'):
        h = doc.add_heading('Abstract', level=1)
        # Some styles want Abstract centered
        doc.add_paragraph(ir['abstract'])

    ## Add body paragraphs
    refs_started = False
    for para in ir['paragraphs']:
        # Skip what we've already rendered
        if para['type'] in ['title', 'abstract']:
            continue
            
        if para['type'] == 'heading':
            level = para.get('heading_level', 1)
            h = doc.add_heading(para['text'], level=min(level, 4))
            if para.get('align') == 'center':
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        elif para['type'] == 'reference' and not refs_started:
            # We handle the reference list at the end
            refs_started = True
            continue
            
        elif para['type'] == 'body' and para['text'].strip():
            p = doc.add_paragraph(para['text'])
            # Support basic formatting if captured in Paragraphs
            if para.get('is_bold'):
                if p.runs:
                    p.runs[0].bold = True
                else:
                    p.add_run().bold = True

    ## Add formatted reference list
    if ir.get('references_formatted') or ir.get('references_raw'):
        doc.add_page_break()
        ref_title = 'Works Cited' if ir.get('style_applied') == 'MLA 9th' else 'References'
        h = doc.add_heading(ref_title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        refs_to_render = ir.get('references_formatted') or ir.get('references_raw', [])
        for ref in refs_to_render:
            p = doc.add_paragraph(ref)
            # Apply hanging indent for professional academic look
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.save(output_path)
    return output_path


# --- Flask Routes ---
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return redirect(url_for('index'))
    
    file = request.files['file']
    style = request.form.get('style', 'APA 7th Edition')
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        job_id = str(uuid.uuid4())
        
        # Ensure directories exist
        Path(app.config['UPLOAD_FOLDER']).mkdir(exist_ok=True)
        Path(app.config['OUTPUT_FOLDER']).mkdir(exist_ok=True)
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_{filename}")
        file.save(filepath)
        
        # Initialize job in DB
        save_job(job_id, filename, style)
        
        try:
            # 2. ingest_document and 3. detect_structure
            ir = ingest_document(filepath, job_id)
            ir = detect_structure(ir)
            
            # Additional metadata for Phase 3
            ir['job_id'] = job_id
            style = request.form.get('style', 'APA 7th Edition')
            
            # 4. Save ir as JSON for debugging
            ir_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_ir.json")
            with open(ir_path, 'w', encoding='utf-8') as f:
                json.dump(ir, f, indent=2)
            
            # 5. Store in-memory
            JOBS_IR[job_id] = ir
            
            print(f"DEBUG: Job {job_id} - Detected Style: {ir['detected_style']}")

            # 6. Formatting Engine (Phase 3)
            ir = format_document(ir, style)
            
            validation = validate_document(ir)
            ir['validation'] = validation
            
            # Render final output
            output_path = render_docx(ir, job_id)
            
            # Update job as complete
            update_job(job_id, 'completed', score=validation['score'], 
                       changes=len(ir.get('change_log', [])))
            
        except Exception as e:
            import traceback
            print(f"Pipeline Error: {e}")
            print(traceback.format_exc())
            update_job(job_id, 'failed', error=str(e))
            
        return redirect(url_for('result', job_id=job_id))
    
    return redirect(url_for('index'))

@app.route('/result/<job_id>')
def result(job_id):
    job = get_job(job_id)
    if not job:
        return "Job not found", 404
    return render_template('result.html', job=job)

@app.route('/download/<job_id>')
def download(job_id):
    job = get_job(job_id)
    if not job:
        return "Job not found", 404
        
    output_filename = f"{job_id}.docx"
    output_path = os.path.abspath(os.path.join(app.config['OUTPUT_FOLDER'], output_filename))
    
    if not os.path.exists(output_path):
        return "File not found", 404
        
    return send_file(output_path, as_attachment=True, download_name=f"formatted_{job['filename']}.docx")

@app.route('/health')
def health():
    return jsonify({'status': 'ok', 'version': '1.0'})

# --- Main Block ---
if __name__ == '__main__':
    # Setup directories
    Path(app.config['UPLOAD_FOLDER']).mkdir(exist_ok=True)
    Path(app.config['OUTPUT_FOLDER']).mkdir(exist_ok=True)
    
    # Initialize database
    init_db()
    
    # Run application
    app.run(debug=True, port=5000)
